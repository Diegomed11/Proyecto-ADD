"""
Exportación del reporte automático a documentos descargables (DOCX y PDF).

Recibe el diccionario que produce el endpoint /api/report y genera un
documento con las mismas secciones: calidad, hallazgos, correlaciones,
limpieza sugerida y modelado.

  to_docx(report) -> bytes   (Word, vía python-docx)
  to_pdf(report)  -> bytes   (PDF, vía fpdf2)
"""

from io import BytesIO
from typing import Any, Dict


# ─────────────────────────── helpers comunes ────────────────────────────────

def _tipo_modelado(tipo: str) -> str:
    return {
        "clasificacion": "Clasificación",
        "regresion": "Regresión",
        "clustering": "Agrupamiento (no supervisado)",
    }.get(tipo, tipo or "—")


# ───────────────────────────────── DOCX ─────────────────────────────────────

def to_docx(report: Dict[str, Any]) -> bytes:
    """Genera el informe en formato Word (.docx) y lo devuelve como bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    source = report.get("source", "fuente")
    dims = report.get("dimensions", {})

    # Título
    h = doc.add_heading("Informe automático de datos", level=0)
    p = doc.add_paragraph()
    run = p.add_run(f"Fuente: {source}  ·  {dims.get('rows', 0)} filas × {dims.get('cols', 0)} columnas")
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 1. Calidad
    cal = report.get("calidad", {})
    doc.add_heading("Calidad de datos", level=1)
    pc = doc.add_paragraph()
    r1 = pc.add_run(f"{cal.get('score', 0)}/100")
    r1.bold = True
    r1.font.size = Pt(22)
    pc.add_run(f"   ·  nota {cal.get('grade', '—')}").bold = True
    doc.add_paragraph(cal.get("resumen", ""))
    for f in cal.get("factores", []):
        doc.add_paragraph(
            f"{f.get('nombre', '')} (−{f.get('penalizacion', 0)}): {f.get('detalle', '')}",
            style="List Bullet",
        )

    # 2. Hallazgos
    doc.add_heading("Hallazgos clave", level=1)
    obs = report.get("observaciones", [])
    if obs:
        for o in obs:
            doc.add_paragraph(o.get("texto", ""), style="List Bullet")
    else:
        doc.add_paragraph("Sin observaciones.")

    # 3. Correlaciones
    doc.add_heading("Correlaciones más fuertes", level=1)
    corr = report.get("correlaciones_top", [])
    if corr:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Variable A", "Variable B", "Correlación"
        for c in corr:
            row = t.add_row().cells
            row[0].text = str(c.get("a", ""))
            row[1].text = str(c.get("b", ""))
            row[2].text = f"{c.get('valor', 0):.2f}"
    else:
        doc.add_paragraph("Sin correlaciones numéricas suficientes.")

    # 4. Limpieza sugerida
    doc.add_heading("Limpieza sugerida", level=1)
    lim = report.get("limpieza_sugerida", {})
    antes, despues = lim.get("antes", {}), lim.get("despues", {})
    doc.add_paragraph(
        f"Filas {antes.get('filas', 0)} → {despues.get('filas', 0)}  ·  "
        f"Columnas {antes.get('columnas', 0)} → {despues.get('columnas', 0)}  ·  "
        f"Nulos {antes.get('nulos', 0)} → {despues.get('nulos', 0)}  ·  "
        f"Duplicados {antes.get('duplicados', 0)} → {despues.get('duplicados', 0)}"
    )
    log = lim.get("log", [])
    if log:
        for e in log:
            doc.add_paragraph(f"{e.get('paso', '')}: {e.get('detalle', '')}", style="List Bullet")
    else:
        doc.add_paragraph("Los datos ya están limpios.")

    # 5. Modelado
    doc.add_heading("Modelado automático", level=1)
    mod = report.get("modelado", {})
    if mod.get("disponible"):
        meta = _tipo_modelado(mod.get("tipo", ""))
        linea = f"Tipo: {meta}"
        if mod.get("objetivo"):
            linea += f"  ·  Objetivo: {mod['objetivo']}"
        doc.add_paragraph(linea)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Modelo", "Métrica", "Valor"
        for c in mod.get("comparacion", []):
            row = t.add_row().cells
            nombre = c.get("nombre", "")
            if nombre == mod.get("recomendado"):
                nombre += "  (recomendado)"
            row[0].text = nombre
            row[1].text = str(c.get("metrica", ""))
            row[2].text = str(c.get("valor", ""))
    else:
        doc.add_paragraph(mod.get("motivo") or "No disponible.")

    foot = doc.add_paragraph()
    fr = foot.add_run("Generado por APCD — Pipeline de Análisis de Datos.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ────────────────────────────────── PDF ─────────────────────────────────────

# Las fuentes núcleo de PDF son latin-1; sustituimos puntuación Unicode común.
_PDF_REPL = {
    "—": "-", "–": "-", "«": '"', "»": '"', "“": '"', "”": '"',
    "‘": "'", "’": "'", "…": "...", "→": "->", "✓": "-", "•": "*", "−": "-",
}


def _s(text: Any) -> str:
    """Hace un texto seguro para las fuentes núcleo de fpdf (latin-1)."""
    s = str(text)
    for a, b in _PDF_REPL.items():
        s = s.replace(a, b)
    return s.encode("latin-1", "replace").decode("latin-1")


def to_pdf(report: Dict[str, Any]) -> bytes:
    """Genera el informe en formato PDF y lo devuelve como bytes."""
    from fpdf import FPDF

    AMBER = (217, 119, 6)
    GRAY = (102, 102, 102)
    DARK = (26, 26, 26)

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    ancho = pdf.w - 2 * pdf.l_margin

    def titulo_seccion(txt):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 8, _s(txt), new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(*AMBER)
        pdf.set_line_width(0.5)
        y = pdf.get_y()
        pdf.line(pdf.l_margin, y, pdf.l_margin + ancho, y)
        pdf.ln(2)

    def parrafo(txt, size=11, color=DARK, bold=False):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.set_text_color(*color)
        pdf.multi_cell(ancho, 5.5, _s(txt))

    def vinheta(txt):
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(ancho, 5, _s("- " + txt))

    source = report.get("source", "fuente")
    dims = report.get("dimensions", {})

    # Título
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 10, _s("Informe automatico de datos"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 10)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 6, _s(f"Fuente: {source}  -  {dims.get('rows', 0)} filas x {dims.get('cols', 0)} columnas"),
             new_x="LMARGIN", new_y="NEXT")

    # 1. Calidad
    cal = report.get("calidad", {})
    titulo_seccion("Calidad de datos")
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(*AMBER)
    pdf.cell(0, 12, _s(f"{cal.get('score', 0)}/100  -  nota {cal.get('grade', '-')}"),
             new_x="LMARGIN", new_y="NEXT")
    parrafo(cal.get("resumen", ""))
    for f in cal.get("factores", []):
        vinheta(f"{f.get('nombre', '')} (-{f.get('penalizacion', 0)}): {f.get('detalle', '')}")

    # 2. Hallazgos
    titulo_seccion("Hallazgos clave")
    obs = report.get("observaciones", [])
    if obs:
        for o in obs:
            vinheta(o.get("texto", ""))
    else:
        parrafo("Sin observaciones.")

    # 3. Correlaciones
    titulo_seccion("Correlaciones mas fuertes")
    corr = report.get("correlaciones_top", [])
    if corr:
        for c in corr:
            vinheta(f"{c.get('a', '')}  <->  {c.get('b', '')}:  {c.get('valor', 0):.2f}")
    else:
        parrafo("Sin correlaciones numericas suficientes.")

    # 4. Limpieza
    titulo_seccion("Limpieza sugerida")
    lim = report.get("limpieza_sugerida", {})
    antes, despues = lim.get("antes", {}), lim.get("despues", {})
    parrafo(
        f"Filas {antes.get('filas', 0)} -> {despues.get('filas', 0)}   "
        f"Columnas {antes.get('columnas', 0)} -> {despues.get('columnas', 0)}   "
        f"Nulos {antes.get('nulos', 0)} -> {despues.get('nulos', 0)}   "
        f"Duplicados {antes.get('duplicados', 0)} -> {despues.get('duplicados', 0)}",
        size=10, color=GRAY,
    )
    log = lim.get("log", [])
    if log:
        for e in log:
            vinheta(f"{e.get('paso', '')}: {e.get('detalle', '')}")
    else:
        parrafo("Los datos ya estan limpios.")

    # 5. Modelado
    titulo_seccion("Modelado automatico")
    mod = report.get("modelado", {})
    if mod.get("disponible"):
        meta = _tipo_modelado(mod.get("tipo", ""))
        linea = f"Tipo: {meta}"
        if mod.get("objetivo"):
            linea += f"   Objetivo: {mod['objetivo']}"
        parrafo(linea, size=10, color=GRAY)
        for c in mod.get("comparacion", []):
            nombre = c.get("nombre", "")
            rec = "  (RECOMENDADO)" if nombre == mod.get("recomendado") else ""
            parrafo(f"{nombre}{rec}  ->  {c.get('metrica', '')}: {c.get('valor', '')}", bold=True)
            parrafo(c.get("interpretacion", ""), size=9, color=GRAY)
    else:
        parrafo(mod.get("motivo") or "No disponible.")

    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(153, 153, 153)
    pdf.cell(0, 5, _s("Generado por APCD - Pipeline de Analisis de Datos."),
             align="C", new_x="LMARGIN", new_y="NEXT")

    salida = pdf.output()
    return bytes(salida)
